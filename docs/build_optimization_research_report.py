"""Build a report-ready DOCX explaining the complete optimization journey."""

from __future__ import annotations

import csv
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_project_description import (
    BLUE,
    CALLOUT,
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT_BLUE,
    MUTED,
    NAVY,
    RED,
    add_body,
    add_bullets,
    add_callout,
    add_code,
    add_figure,
    add_heading,
    add_numbers,
    add_page_number,
    add_table as base_add_table,
    configure_page,
    configure_styles,
    set_run_font,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "project"
ASSETS = PROJECT / "docs" / "_optimization_research_assets"
OUTPUT = ROOT / "项目优化全过程深度研究与论文写作参考.docx"

SCHOOL_FULL = PROJECT / "results" / "experiments" / \
    "20260718-152329241-medium-fast-normalization-full" / "summary.csv"
SCHOOL_TOP50 = PROJECT / "results" / "experiments" / \
    "20260718-152600213-medium-fast-normalization-top50" / "summary.csv"
SCHOOL_LARGE = PROJECT / "results" / "experiments" / \
    "20260718-154737913-large-fast-normalization-top50" / "summary.csv"
BASELINE_MEDIUM = ROOT / "results" / "summary" / \
    "runtime-medium-20260716-125503-summary.csv"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def serial_row(rows: list[dict[str, str]]) -> dict[str, str]:
    return next(row for row in rows if row["version"] == "serial")


def omp_rows(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    return sorted((row for row in rows if row["version"] == "openmp"),
                  key=lambda row: int(row["threads"]))


def add_table(doc, headers, rows, widths, font_size=9.0, header_fill=LIGHT_BLUE):
    table = base_add_table(doc, headers, rows, widths, font_size, header_fill)
    set_table_geometry(table, widths)
    return table


def source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Figure Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    set_run_font(paragraph.add_run(f"数据来源：{text}"), size=8.5,
                 italic=True, color=MUTED)


def set_custom_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    for run in list(header_p.runs):
        header_p._p.remove(run._r)
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header_p.add_run("Instacart 购物篮推荐  |  优化全过程研究"),
                 size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    for run in list(footer_p.runs):
        footer_p._p.remove(run._r)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer_p.add_run("第 "), size=8.5, color=MUTED)
    add_page_number(footer_p)
    set_run_font(footer_p.add_run(" 页  |  论文报告组参考稿"),
                 size=8.5, color=MUTED)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("并行程序设计课程大作业 · 技术研究报告"),
                 size=11.5, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    set_run_font(title.add_run("Instacart 购物篮推荐项目\n优化全过程深度研究"),
                 size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    set_run_font(subtitle.add_run(
        "从问题建模、串并行基线到瓶颈迁移、并行邻接构建与分桶归并"),
        size=13, color=DARK_BLUE)

    add_table(doc, ["文档属性", "说明"], [
        ["用途", "供论文报告组提炼项目背景、算法、实验设计、性能分析与优化思路"],
        ["技术主线", "C11 + OpenMP；稀疏共现图；用户级 Top-K 推荐"],
        ["证据范围", "toy/small 正确性、medium/large 性能、1-48 线程学校平台实验"],
        ["当前结论", "并行邻接构建已验证；分桶并行归并为下一阶段核心策略"],
        ["编制日期", "2026 年 7 月 18 日"],
    ], [1.35, 5.15], font_size=9.8)

    add_callout(doc, "核心叙事",
                "本项目最有价值的部分不是某一个最终加速比，而是通过分阶段计时不断发现瓶颈迁移：先并行共现与推荐，再优化算法本身，随后解释加速比下降，最后继续并行化新的主导阶段。",
                color=GOLD, fill="FFF8E8")
    doc.add_page_break()


def chart_style() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"
    plt.rcParams["axes.edgecolor"] = "#B7C1CC"
    plt.rcParams["grid.color"] = "#D9E0E7"
    plt.rcParams["grid.alpha"] = 0.65


def save_figure(path: Path) -> None:
    plt.tight_layout()
    plt.savefig(path, dpi=190, bbox_inches="tight")
    plt.close()


def create_workflow_chart(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.4, 5.4))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)
    ax.axis("off")
    boxes = [
        (0.25, 3.75, 1.75, 1.25, "原始 CSV\n订单/商品", "#E8EEF5"),
        (2.25, 3.75, 1.75, 1.25, "CSR 数据\n用户历史", "#DDEBF7"),
        (4.25, 3.75, 1.75, 1.25, "共现统计\nPairHashMap", "#D9EAD3"),
        (6.25, 3.75, 1.75, 1.25, "CSR 邻接图\nTop-N", "#FFF2CC"),
        (8.25, 3.75, 1.75, 1.25, "用户候选\n评分与 Top-K", "#FCE5CD"),
        (10.25, 3.75, 1.5, 1.25, "评估\n校验和", "#EADCF8"),
        (2.1, 1.0, 2.25, 1.2, "订单级 OpenMP\n线程局部哈希", "#EAF4E5"),
        (4.85, 1.0, 2.25, 1.2, "并行邻接构建\n两遍式无锁写入", "#EAF4E5"),
        (7.6, 1.0, 2.25, 1.2, "用户级 OpenMP\n线程私有工作区", "#EAF4E5"),
    ]
    for x, y, w, h, label, color in boxes:
        patch = FancyBboxPatch((x, y), w, h,
                               boxstyle="round,pad=0.04,rounding_size=0.12",
                               linewidth=1.4, edgecolor="#607284", facecolor=color)
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=11, color="#243447", fontweight="bold")
    for x1, x2 in ((2.0, 2.25), (4.0, 4.25), (6.0, 6.25),
                   (8.0, 8.25), (10.0, 10.25)):
        ax.add_patch(FancyArrowPatch((x1, 4.38), (x2, 4.38), arrowstyle="-|>",
                                     mutation_scale=14, color="#657786", lw=1.6))
    for x in (3.2, 5.95, 8.7):
        ax.add_patch(FancyArrowPatch((x, 3.72), (x, 2.22), arrowstyle="<->",
                                     mutation_scale=13, color="#2F6B4F", lw=1.5))
    ax.text(6, 5.55, "端到端算法数据流与三类并行点", ha="center",
            fontsize=17, fontweight="bold", color="#183B56")
    save_figure(path)


def create_baseline_scaling(path: Path) -> None:
    rows = read_csv(BASELINE_MEDIUM)
    omp = omp_rows(rows)
    threads = [int(row["threads"]) for row in omp]
    speedup = [float(row["speedup"]) for row in omp]
    efficiency = [100 * float(row["efficiency"]) for row in omp]
    fig, ax1 = plt.subplots(figsize=(8.8, 4.9))
    x = range(len(threads))
    ax1.plot(x, speedup, marker="o", lw=2.4, color="#2E74B5", label="加速比")
    ax1.plot(x, threads, ls="--", lw=1.5, color="#9AA7B4", label="理想加速")
    ax1.set_xticks(list(x), threads)
    ax1.set_xlabel("OpenMP 线程数")
    ax1.set_ylabel("Speedup")
    ax1.grid(True, axis="y")
    ax2 = ax1.twinx()
    ax2.plot(x, efficiency, marker="s", lw=2, color="#9A6A00", label="并行效率")
    ax2.set_ylabel("Efficiency (%)")
    lines = ax1.get_lines() + ax2.get_lines()
    ax1.legend(lines, [line.get_label() for line in lines], loc="upper left")
    ax1.set_title("基础 Full 方案：medium 串并行扩展性")
    save_figure(path)


def create_quality_tradeoff(path: Path) -> None:
    labels = ["Full", "Top-20", "Top-50", "Top-100"]
    runtime = [1431, 710, 748, 826]
    hit = [0.8428125, 0.8403125, 0.84125, 0.8396875]
    recall = [0.31472267698, 0.30534367814, 0.308845818812,
              0.310919533357]
    fig, ax1 = plt.subplots(figsize=(8.8, 5.0))
    x = list(range(len(labels)))
    bars = ax1.bar(x, runtime, width=0.55, color="#A9C7E5", label="8线程算法时间")
    ax1.set_xticks(x, labels)
    ax1.set_ylabel("algorithm_ms")
    ax1.set_ylim(0, 1600)
    ax1.bar_label(bars, fmt="%.0f", padding=3, fontsize=9)
    ax2 = ax1.twinx()
    ax2.plot(x, hit, marker="o", lw=2.2, color="#2F6B4F", label="Hit Rate@10")
    ax2.plot(x, recall, marker="s", lw=2.2, color="#9A6A00", label="Recall@10")
    ax2.set_ylim(0.28, 0.86)
    ax2.set_ylabel("推荐质量")
    lines = [bars] + ax2.get_lines()
    ax1.legend(lines, [item.get_label() for item in lines], loc="upper right")
    ax1.set_title("Top-N 截断的性能—质量权衡（medium，8线程）")
    save_figure(path)


def create_speedup_journey(path: Path) -> None:
    labels = ["基础 Full\n8线程", "Top-50\n邻接仍串行", "Top-50\n并行邻接", "学校平台\nTop-50 24线程"]
    values = [4.915, 1.310, 3.318, 3.739]
    colors = ["#2E74B5", "#C98983", "#5B9A72", "#1F6F8B"]
    fig, ax = plt.subplots(figsize=(9.2, 4.9))
    bars = ax.bar(range(len(labels)), values, color=colors, width=0.62)
    ax.set_xticks(range(len(labels)), labels)
    ax.set_ylabel("同算法串行/OpenMP 加速比")
    ax.set_ylim(0, 5.5)
    ax.grid(True, axis="y")
    ax.bar_label(bars, labels=[f"{value:.3f}×" for value in values],
                 padding=4, fontsize=10)
    ax.set_title("优化全过程中的加速比变化：下降不是失败，而是瓶颈迁移")
    save_figure(path)


def create_school_scaling(path: Path) -> None:
    series = [
        ("medium Full", read_csv(SCHOOL_FULL), "#2E74B5", "o"),
        ("medium Top-50", read_csv(SCHOOL_TOP50), "#2F6B4F", "s"),
        ("large Top-50", read_csv(SCHOOL_LARGE), "#9A6A00", "^"),
    ]
    fig, ax = plt.subplots(figsize=(9.3, 5.2))
    positions = None
    thread_labels = None
    for label, rows, color, marker in series:
        omp = omp_rows(rows)
        threads = [int(row["threads"]) for row in omp]
        speedups = [float(row["speedup"]) for row in omp]
        if positions is None:
            positions = list(range(len(threads)))
            thread_labels = threads
        ax.plot(positions, speedups, marker=marker, lw=2.2,
                color=color, label=label)
    ax.set_xticks(positions, thread_labels)
    ax.set_xlabel("线程数")
    ax.set_ylabel("相对独立串行基线的加速比")
    ax.grid(True, axis="y")
    ax.legend()
    ax.set_title("学校大数据平台 1-48 线程扩展性")
    save_figure(path)


def create_large_stage_chart(path: Path) -> None:
    rows = read_csv(SCHOOL_LARGE)
    omp = omp_rows(rows)
    threads = [int(row["threads"]) for row in omp]
    fields = [
        ("median_cooccur_compute_ms", "共现计算", "#7FAED7"),
        ("median_merge_ms", "局部表归并", "#D98780"),
        ("median_adjacency_ms", "邻接构建", "#89B68D"),
        ("median_recommend_ms", "推荐", "#D6B46A"),
    ]
    x = list(range(len(threads)))
    bottom = [0.0] * len(threads)
    fig, ax = plt.subplots(figsize=(9.5, 5.5))
    for field, label, color in fields:
        values = [float(row[field]) for row in omp]
        ax.bar(x, values, bottom=bottom, label=label, color=color, width=0.68)
        bottom = [a + b for a, b in zip(bottom, values)]
    ax.set_xticks(x, threads)
    ax.set_xlabel("线程数")
    ax.set_ylabel("阶段中位数（ms）")
    ax.legend(ncol=2)
    ax.set_title("large Top-50：线程增加后的阶段瓶颈迁移")
    ax.grid(True, axis="y")
    save_figure(path)


def create_merge_share_chart(path: Path) -> None:
    rows = read_csv(SCHOOL_LARGE)
    omp = omp_rows(rows)
    threads = [int(row["threads"]) for row in omp]
    shares = [100 * float(row["median_merge_ms"]) /
              float(row["median_algorithm_ms"]) for row in omp]
    runtimes = [float(row["median_algorithm_ms"]) for row in omp]
    fig, ax1 = plt.subplots(figsize=(9.1, 5.0))
    x = list(range(len(threads)))
    ax1.plot(x, shares, marker="o", lw=2.4, color="#A23B3B", label="归并占比")
    ax1.fill_between(x, shares, alpha=0.12, color="#A23B3B")
    ax1.set_xticks(x, threads)
    ax1.set_xlabel("线程数")
    ax1.set_ylabel("merge_ms / algorithm_ms (%)")
    ax1.set_ylim(0, 75)
    ax1.grid(True, axis="y")
    ax2 = ax1.twinx()
    ax2.plot(x, runtimes, marker="s", lw=2, color="#2E74B5", label="算法时间")
    ax2.set_ylabel("algorithm_ms")
    lines = ax1.get_lines() + ax2.get_lines()
    ax1.legend(lines, [line.get_label() for line in lines], loc="center right")
    ax1.set_title("large Top-50：归并占比上升导致 24 线程后反向变慢")
    save_figure(path)


def create_bucket_forecast(path: Path) -> None:
    serial = 34974.352
    current_total = 9407.263
    merge = 5553.249
    non_merge = current_total - merge
    factors = [1, 2, 4, 8, 10_000]
    labels = ["当前", "归并2×", "归并4×", "归并8×", "理想上限"]
    totals = [non_merge + (merge / factor if factor < 100 else 0)
              for factor in factors]
    speedups = [serial / total for total in totals]
    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    bars = ax.bar(range(len(labels)), speedups,
                  color=["#A9C7E5", "#8CB3D9", "#6E9DCB", "#4D83B8", "#2F6B4F"])
    ax.set_xticks(range(len(labels)), labels)
    ax.set_ylabel("预计总体加速比")
    ax.set_ylim(0, 10)
    ax.grid(True, axis="y")
    ax.bar_label(bars, labels=[f"{value:.2f}×" for value in speedups],
                 padding=4)
    ax.set_title("基于 large 24线程实测阶段时间的分桶归并收益上界")
    save_figure(path)


def create_figures() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    chart_style()
    figures = {
        "workflow": ASSETS / "01-workflow.png",
        "baseline": ASSETS / "02-baseline-scaling.png",
        "quality": ASSETS / "03-quality-tradeoff.png",
        "journey": ASSETS / "04-speedup-journey.png",
        "school": ASSETS / "05-school-scaling.png",
        "stages": ASSETS / "06-large-stages.png",
        "merge": ASSETS / "07-merge-share.png",
        "bucket": ASSETS / "08-bucket-forecast.png",
    }
    create_workflow_chart(figures["workflow"])
    create_baseline_scaling(figures["baseline"])
    create_quality_tradeoff(figures["quality"])
    create_speedup_journey(figures["journey"])
    create_school_scaling(figures["school"])
    create_large_stage_chart(figures["stages"])
    create_merge_share_chart(figures["merge"])
    create_bucket_forecast(figures["bucket"])
    return figures


def add_navigation(doc: Document) -> None:
    add_heading(doc, "内容导航", 1)
    add_body(doc, "本文按照项目真实演进顺序组织，而不是按最终代码倒序解释。报告组可以直接把以下章节映射到课程论文。")
    add_table(doc, ["部分", "核心内容", "可用于论文"], [
        ["第一部分", "问题定义、数据、数学建模", "问题描述、算法原理"],
        ["第二部分", "串行基线、基础 OpenMP 与审效体系", "程序设计、并行设计"],
        ["第三部分", "哈希归并异常、分母预计算、Top-N", "优化过程、质量权衡"],
        ["第四部分", "加速比下降与并行邻接构建", "Amdahl 定律、瓶颈迁移"],
        ["第五部分", "1-48线程结果与分桶归并", "实验结果、扩展性分析"],
        ["附录", "指标口径、证据路径、写作建议", "复现实验、论文素材"],
    ], [1.0, 3.25, 2.25], font_size=9.3)
    add_callout(doc, "阅读提示",
                "不同日期、不同机器的绝对时间不可直接横向拼接。本文跨阶段主要比较同一实验批次内部的串并行加速比；跨批次数据仅用于解释优化思路和瓶颈变化。",
                color=RED, fill="FCEEEE")


def add_executive_summary(doc: Document) -> None:
    add_heading(doc, "研究摘要：项目真正优化了什么", 1, page_break=True)
    add_body(doc, "项目将 Instacart 下一购物篮预测转化为基于历史复购和商品共现的 Top-K 推荐问题。选择该模型并非为了追求 Kaggle 排名，而是因为订单之间天然独立、单个购物篮两两组合计算密集、用户推荐也能独立并行，能够完整展示 OpenMP 中的任务划分、线程私有数据、负载均衡、归并开销和 Amdahl 定律。")
    add_body(doc, "整个优化过程经历了三次瓶颈迁移。第一阶段，完整共现图下推荐计算占据主要时间，订单级共现和用户级推荐获得了明显加速；第二阶段，精确分母预计算与 Top-N 截断大幅降低推荐时间，却使串行构图和归并占比上升，同算法加速比反而变小；第三阶段，两遍式并行邻接构建解决了构图瓶颈，学校平台 1-48 线程实验进一步证明串行归并已成为新主导阶段。")
    add_callout(doc, "最终证据",
                "学校平台上 medium Full 在 48 线程达到 4.936 倍；medium Top-50 在 24 线程达到 3.739 倍；large Top-50 在 24 线程达到 3.718 倍。large 48线程中归并耗时 7263 ms，占算法时间约 66.8%，因此下一项优化应是分桶并行归并，而不是盲目增加推荐模型复杂度。",
                color=GREEN, fill="EAF4E5")
    add_bullets(doc, [
        "算法质量：medium Top-50 的 Hit Rate@10 仅比 Full 下降 0.001563，Recall@10 下降 0.005877，低于 0.01 质量门。",
        "工程正确性：toy/small 完整逐项比较，medium/large 使用稳定校验和；串行/OpenMP 在所有正式线程点上保持一致。",
        "性能方法：所有正式结果采用 Release、预热、重复运行和中位数；计时区分 I/O、共现计算、归并、邻接构建、推荐和评估。",
        "研究价值：优化不是单向追求更快，而是通过实验说明算法优化会改变并行比例，并驱动下一轮并行设计。",
    ])


def add_problem_model(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "第一部分  从业务问题到可并行计算模型", 1, page_break=True)
    add_heading(doc, "1. 原始问题与课程化建模", 2)
    add_body(doc, "原始 Instacart 任务要求根据用户历史订单预测下一购物篮中的商品集合。项目将其简化为可解释的共现推荐：每个订单视为购物篮，每种商品视为图节点，同一订单中任意两种商品形成无向边，共现次数作为边权；用户历史商品作为查询起点，聚合其邻居并输出 Top-K。")
    add_body(doc, "这一建模保留了三个实际信号：用户个人复购、商品搭配关系和全局流行度；同时避开模型训练、参数搜索和分布式框架，把课程重点集中在稀疏数据结构、OpenMP 并行和性能分析。")
    add_code(doc, "score(u,p) = 1.0 * freq(u,p)\n"
                  "           + 0.8 * co_score(u,p)\n"
                  "           + 0.2 * log(1 + popularity(p))\n\n"
                  "co_score(u,p) = Σ freq(u,q) * cooccur(q,p)\n"
                  "                  / sqrt(popularity(q)*popularity(p))")
    add_body(doc, "评分相同时按商品 ID 升序，保证串行和并行版本不受哈希遍历顺序影响。该确定性规则后来成为所有 checksum 和逐项比较能够成立的基础。")

    add_heading(doc, "2. 数据处理与规模分层", 2)
    add_table(doc, ["规模", "订单数", "prior 明细", "train 明细", "商品数", "职责"], [
        ["toy", "7", "6", "4", "4", "手算预言机与边界测试"],
        ["small", "2,995", "26,575", "1,327", "6,170", "完整正确性回归"],
        ["medium", "81,832", "761,750", "33,782", "29,330", "重复性能与参数审效"],
        ["large", "3,421,083", "32,434,489", "1,384,617", "49,685", "规模与高线程验证"],
    ], [0.65, 1.0, 1.15, 1.05, 0.9, 1.75], font_size=8.8)
    source_note(doc, "问题分析与数据处理.docx；project/docs/test-report.md")
    add_body(doc, "数据清洗包括去重、首单 days_since_prior_order 缺失处理、ID 整型化和订单引用一致性检查。核心程序读取 orders、prior、train、products 四张表，并将购物篮、用户历史和验证真值转换为 CSR 连续数组。")
    add_bullets(doc, [
        "购物篮 CSR：offsets 定位每个订单的商品区间，避免 vector 嵌套开销。",
        "用户历史 CSR：商品 ID 与购买频次连续存储，支持顺序扫描。",
        "验证真值 CSR：按用户保存 train 订单商品，便于离线评估。",
        "共现 PairHashMap：64 位商品对编码 + 开放寻址，避免字符串 key。",
    ])
    add_figure(doc, figures["workflow"], "图 1  端到端数据流、核心数据结构与三类 OpenMP 并行点", 6.35)

    add_heading(doc, "3. 计算复杂度与并行潜力", 2)
    add_body(doc, "对长度为 m 的购物篮，需要枚举 m(m-1)/2 个无序商品对，因此共现阶段复杂度为 O(Σm_i²)。不同订单互不依赖，但订单长度不同会造成平方级负载差异；用户推荐阶段的工作量又取决于历史商品数和邻居度数。由此形成两个天然并行层次：订单级和用户级。")
    add_callout(doc, "建模取舍",
                "时间衰减、类别亲和度和负采样具有推荐研究价值，但会扩大数据模型和参数空间。基础版本优先保证可解释、可复现和串并行语义一致；Top-N 是唯一进入主线的质量—性能优化。",
                color=BLUE)


def add_serial_parallel(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "第二部分  串行基线与基础 OpenMP 实现", 1, page_break=True)
    add_heading(doc, "4. 从 C++ 串行原型到独立 C11 基线", 2)
    add_body(doc, "串行同学首先在 basket_recommender 目录中完成 C++ 原型：unordered_map 统计共现、构建 Top-N 邻居、按用户生成推荐，并计算 Hit Rate 与 Recall。该版本证明了业务闭环和公式可运行，但字符串商品对、嵌套容器和模块间数据复制不利于 large 规模和精确性能分析。")
    add_body(doc, "并行项目没有把“OpenMP 设置 1 线程”当作串行基线，而是重新实现独立 C11 串行路径。这样 T_serial 不包含线程运行库、线程局部表和归并开销，能够公平反映算法本身。C11 版本进一步使用 uint64_t 商品对编码、开放寻址哈希表和 CSR 连续存储。")
    add_table(doc, ["层次", "串行原型", "C11/OpenMP 主项目"], [
        ["商品对", "字符串 a,b", "64位整数编码"],
        ["稀疏统计", "unordered_map", "开放寻址 PairHashMap"],
        ["历史/购物篮", "嵌套 vector/map", "CSR 连续数组"],
        ["串行基线", "业务参考实现", "独立普通循环，作为 T_serial"],
        ["并行路径", "无", "订单级共现 + 用户级推荐 + 后续并行构图"],
    ], [1.1, 2.25, 3.15], font_size=9.1)

    add_heading(doc, "5. 基础 OpenMP：线程局部统计", 2)
    add_body(doc, "共现阶段采用 dynamic,64 调度购物篮。每个线程维护独立 PairHashMap、商品热度数组和商品对事件数，从根源上避免在热点双重循环中使用 critical 或全局锁。并行区结束后，再把局部表依次归并到全局表。")
    add_code(doc, "#pragma omp parallel\n"
                  "{\n"
                  "    tid = omp_get_thread_num();\n"
                  "    local_pairs = pair_maps[tid];\n"
                  "    #pragma omp for schedule(runtime)\n"
                  "    for each basket:\n"
                  "        update local popularity\n"
                  "        enumerate i < j and update local_pairs\n"
                  "}\n"
                  "merge all thread-local maps")
    add_body(doc, "用户推荐阶段按用户并行。模型、历史和邻接图只读；每个线程拥有独立 RecommendationWorkspace；每个用户写入固定结果区间。因此推荐核心循环不需要锁。动态调度用于缓解不同用户候选量差异。")

    add_heading(doc, "6. 正确性优先的审效体系", 2)
    add_body(doc, "每个优化只有在语义门通过后才能进入性能实验。项目设计了 G0-G10 与 O0-O7 多级质量门：从数据结构和哈希表单元测试，到串并行逐项比较、Top-N 边界、调度确定性、Release 实验格式校验。")
    add_table(doc, ["质量门", "核心检查", "代表性证据"], [
        ["G1-G3", "加载、哈希、串行共现", "toy 手算；small 129,071 唯一边"],
        ["G4", "串并行共现等价", "1/2/4线程，static/dynamic，逐键计数一致"],
        ["G5-G7", "图、推荐、Top-K", "offset单调、无重复、分数容差1e-12"],
        ["O1", "精确分母预计算", "优化前后推荐ID校验和不变"],
        ["O2", "Top-N 截断", "N=0/1/超大；同权重按ID；large质量门"],
        ["O7", "并行邻接构建", "toy/small 1/2/4线程图结构逐字节一致"],
        ["平台门", "1-48线程正式结果", "29条记录/批次，validation.json=pass"],
    ], [0.8, 2.2, 3.5], font_size=8.7)

    add_heading(doc, "7. 基础并行效果", 2)
    add_figure(doc, figures["baseline"], "图 2  基础 Full 方案在 medium 上的加速比与并行效率", 6.15)
    add_body(doc, "基础 Full medium 三次中位数中，串行算法 7711 ms，OpenMP 8线程 1569 ms，加速比 4.915、效率 61.4%。large 单次完整图验证中，串行 1,217,953 ms，8线程 199,149 ms，加速比 6.116。此时推荐阶段仍占主要时间，因此用户级并行能充分发挥。")
    source_note(doc, "results/summary/runtime-medium-20260716-125503-summary.csv；runtime-large-20260716-smoke-summary.csv")


def add_algorithm_optimization(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "第三部分  算法优化与加速比下降", 1, page_break=True)
    add_heading(doc, "8. 第一次异常：OpenMP 1线程反而极慢", 2)
    add_body(doc, "最初 OpenMP 1线程实验中，局部表归并耗时约 12.96 秒，导致总时间 20.49 秒、相对串行加速比仅 0.358。结果虽然正确，但性能显然不合理。分阶段计时帮助排除共现计算和推荐，最终定位到全局开放寻址哈希表从很小容量开始逐步扩容；按局部哈希槽顺序插入时，早期小表形成长线性探测链。")
    add_table(doc, ["状态", "OpenMP1 merge_ms", "algorithm_ms", "处理"], [
        ["异常版本", "12,961", "20,490", "标记 INVALID，不进入论文结论"],
        ["修复版本", "48", "7,885", "按局部键数上界一次 reserve"],
    ], [1.1, 1.6, 1.55, 2.25], font_size=9.2)
    add_callout(doc, "方法论价值",
                "并行结果正确并不代表性能实验有效。阶段计时、异常数据保留和重新跑完整回归，使这个问题从“删掉难看的点”转化为可解释的工程优化案例。",
                color=GOLD, fill="FFF8E8")

    add_heading(doc, "9. 语义保持优化：预计算精确归一化分母", 2)
    add_body(doc, "推荐热点循环原本反复执行 sqrt(popularity(q)*popularity(p))。优化方案在构图时为每条邻接边预计算完整 double 分母，推荐阶段只执行 frequency*weight/denominator。第一次尝试把分母拆成两个逆平方根，因浮点运算顺序变化导致推荐 ID checksum 改变，被审效门撤回；最终方案严格保持原运算顺序。")
    add_table(doc, ["版本", "优化前算法", "优化后算法", "下降", "推荐阶段变化"], [
        ["串行 Full", "7711 ms", "4949 ms", "35.8%", "7054 → 4297 ms"],
        ["OpenMP 8", "1569 ms", "1431 ms", "8.8%", "999 → 818 ms"],
    ], [1.2, 1.2, 1.2, 0.9, 2.0], font_size=9.1)
    add_body(doc, "串行收益更大，因为开方全部集中在单线程推荐热点；并行版本原本已分摊该热点。此处第一次出现重要现象：算法优化越成功，原先可并行阶段占比越低，同算法加速比可能下降。")

    add_heading(doc, "10. Top-N：以小质量损失换取数量级计算削减", 2)
    add_figure(doc, figures["quality"], "图 3  Full、Top-20、Top-50、Top-100 的性能—质量权衡", 6.15)
    add_body(doc, "Top-N 对每个商品仅保留共现权重最高的 N 个邻居。medium 上 Top-20 时间最短，但 large Recall 绝对损失超过 0.01；Top-50 在 large 上通过质量门，因此成为最终实用配置。")
    add_table(doc, ["指标", "Full", "Top-50", "变化"], [
        ["图邻接项", "5,225,290", "1,052,186", "减少79.9%"],
        ["候选总量", "83,335,836", "4,794,337", "减少94.25%"],
        ["Hit Rate@10", "0.842813", "0.841250", "-0.001563"],
        ["Recall@10", "0.314723", "0.308846", "-0.005877"],
        ["NDCG@10", "0.377202", "0.376833", "-0.000370"],
        ["MRR@10", "0.593036", "0.601145", "+0.008109"],
    ], [1.25, 1.55, 1.55, 2.15], font_size=9.0)

    add_heading(doc, "11. 为什么优化后加速比反而更小", 2)
    add_figure(doc, figures["journey"], "图 4  优化过程中同算法加速比的下降、恢复与平台验证", 6.25)
    add_body(doc, "并行邻接实施前，medium Top-50 串行约 956 ms、OpenMP 8线程约 730 ms，同算法加速比只有 1.310。原因不是 OpenMP 推荐失效，而是 Top-N 将推荐阶段从数秒缩短到几十毫秒；串行邻接构建、局部表归并和固定开销成为新的总时间主导。")
    add_code(doc, "Speedup(p) = T_serial / T_parallel(p)\n"
                  "Amdahl: S(p) = 1 / ((1-f) + f/p)\n"
                  "当算法优化使 f 下降时，即使 T_parallel 更短，S(p) 也可能下降。")
    add_callout(doc, "必须分开的三个指标",
                "同算法并行加速比用于证明 OpenMP；Full→Top-50 的算法收益用于证明剪枝有效；Full 串行→Top-50 OpenMP 是组合收益，不能标成纯 OpenMP 加速比。",
                color=RED, fill="FCEEEE")


def add_parallel_graph(doc: Document) -> None:
    add_heading(doc, "第四部分  针对新瓶颈的并行邻接构建", 1, page_break=True)
    add_heading(doc, "12. 从串行构图到两遍式无锁构图", 2)
    add_body(doc, "旧路径先串行遍历全局商品对，统计每个商品度数，再写入双向邻接项，随后逐商品排序和 Top-N。Top-50 后该阶段占串行算法接近一半，因此必须新增独立 cooccur_graph_build_openmp，而不是让串行函数内部悄悄调用 OpenMP。")
    add_numbers(doc, [
        "将 PairHashMap 槽位划分为固定逻辑分区，每个分区扫描自己的槽位范围。",
        "每个分区维护一行 local_degrees[worker][product]，统计边两端度数，不使用 atomic。",
        "按商品归并各分区度数并做前缀和，得到 CSR offsets。",
        "为每个“分区×商品”计算不重叠写入游标，第二遍并行填充双向邻接项。",
        "按商品动态调度 qsort 与 Top-N 复制，使用 max reduction 统计最大度数。",
    ])
    add_code(doc, "worker slot partitions\n"
                  "    ↓ parallel degree count (thread-private rows)\n"
                  "reduce degrees by product → prefix sum → CSR offsets\n"
                  "    ↓ allocate disjoint per-worker write ranges\n"
                  "parallel edge fill → parallel per-product Top-N")

    add_heading(doc, "13. 为什么选择局部度数而不是 atomic", 2)
    add_body(doc, "商品数约 5 万。在 48 线程下，local_degrees 使用 48×50,000×8≈19.2 MB，thread_cursors 再使用约 19.2 MB，总工作区约 38.4 MB，在 large 场景可接受。相比对热门商品度数执行高竞争 atomic，该方案用可预测内存换取无锁写入。")
    add_body(doc, "固定逻辑分区与实际执行线程解耦，即使 OpenMP 运行时调度分区任务，第一次计数和第二次填充仍扫描相同槽位范围。最终邻接段再按商品 ID 排序，因此串行和并行图能够逐字节一致。")

    add_heading(doc, "14. 本地审效结果", 2)
    add_table(doc, ["配置", "串行", "OpenMP8", "加速比", "邻接阶段"], [
        ["medium Top-50", "1732 ms", "522 ms", "3.318×", "887 → 143 ms（6.20×）"],
        ["medium Full", "4801 ms", "1074 ms", "4.470×", "384 → 75 ms（5.12×）"],
    ], [1.45, 1.1, 1.15, 1.0, 2.0], font_size=9.0)
    add_body(doc, "toy 和 small 在 1/2/4 线程下对 offsets、邻居 ID、权重和归一化分母逐字节比较；完整正确性回归、Release 构建和结果验证器自测均通过。由此可以把加速比从 1.31 恢复到 3.318，并确认结果并非通过改变推荐语义获得。")


def school_table_rows(path: Path) -> list[list[str]]:
    rows = read_csv(path)
    serial = serial_row(rows)
    result = [["串行", "1", f"{float(serial['median_algorithm_ms']):.0f}",
               "1.000", "100.0%"]]
    for row in omp_rows(rows):
        result.append(["OpenMP", row["threads"],
                       f"{float(row['median_algorithm_ms']):.0f}",
                       f"{float(row['speedup']):.3f}",
                       f"{100*float(row['efficiency']):.1f}%"])
    return result


def add_school_results(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "第五部分  学校平台 1-48线程验证", 1, page_break=True)
    add_heading(doc, "15. 实验环境与方法", 2)
    add_body(doc, "学校平台为 Linux，检测到 48 个逻辑处理器。实验使用 OMP_PROC_BIND=spread、OMP_PLACES=cores，线程点为 1/2/4/8/12/16/24/32/48；每个 OpenMP 点运行3次，串行2次，预热1次，线程执行顺序随机化。三个正式批次各29条记录，validation.json 均为 pass。")
    add_bullets(doc, [
        "medium Full：观察完整图下纯并行扩展性。",
        "medium Top-50：观察优化后实际配置与瓶颈迁移。",
        "large Top-50：验证大数据、更多用户和40,749,010唯一商品对下的可扩展性。",
        "核心加速比使用 algorithm_ms；end_to_end_ms 另行说明串行 CSV I/O。",
    ])
    add_figure(doc, figures["school"], "图 5  学校平台 medium/large 的 1-48线程加速比", 6.3)

    add_heading(doc, "16. 三组正式结果", 2)
    add_heading(doc, "16.1 medium Full", 3)
    add_table(doc, ["版本", "线程", "algorithm_ms", "Speedup", "Efficiency"],
              school_table_rows(SCHOOL_FULL), [1.0, 0.85, 1.55, 1.25, 1.85], 8.7)
    source_note(doc, "project/results/experiments/20260718-152329241-medium-fast-normalization-full/summary.csv")
    add_body(doc, "最短时间出现在48线程：1222 ms，相对独立串行6032 ms，加速比4.936。24线程为1348 ms、4.475倍，虽然绝对时间略慢，但效率18.6%高于48线程的10.3%。")

    add_heading(doc, "16.2 medium Top-50", 3)
    add_table(doc, ["版本", "线程", "algorithm_ms", "Speedup", "Efficiency"],
              school_table_rows(SCHOOL_TOP50), [1.0, 0.85, 1.55, 1.25, 1.85], 8.7)
    source_note(doc, "project/results/experiments/20260718-152600213-medium-fast-normalization-top50/summary.csv")
    add_body(doc, "24线程达到最佳414 ms、3.739倍；32线程418 ms与24线程基本持平；48线程449 ms反而慢8.4%。这表明 Top-50 在24线程后已经受到串行归并和硬件拓扑限制。")

    add_heading(doc, "16.3 large Top-50", 3)
    add_table(doc, ["版本", "线程", "algorithm_ms", "Speedup", "Efficiency"],
              school_table_rows(SCHOOL_LARGE), [1.0, 0.85, 1.55, 1.25, 1.85], 8.7)
    source_note(doc, "project/results/experiments/20260718-154737913-large-fast-normalization-top50/summary.csv")
    add_body(doc, "large 在24线程达到最佳9407 ms、3.718倍；32线程9982 ms，48线程10876 ms。24→48线程反而慢15.6%。该下降在三次重复中稳定出现，不能归因于单个异常值。")

    add_heading(doc, "17. 高线程瓶颈：归并反向膨胀", 2)
    add_figure(doc, figures["stages"], "图 6  large Top-50 各阶段随线程数变化的堆叠时间", 6.3)
    add_figure(doc, figures["merge"], "图 7  large Top-50 归并占比与总算法时间", 6.2)
    add_table(doc, ["线程", "共现计算", "归并", "邻接", "推荐", "归并占比"], [
        ["1", "10130", "1176", "16315", "7818", "3.3%"],
        ["8", "3791", "3719", "2286", "2809", "29.6%"],
        ["16", "2197", "4929", "1180", "2015", "47.7%"],
        ["24", "1743", "5553", "829", "1378", "59.0%"],
        ["32", "1663", "6318", "843", "1037", "63.3%"],
        ["48", "1712", "7263", "1081", "891", "66.8%"],
    ], [0.65, 1.1, 1.05, 1.0, 1.0, 1.7], font_size=8.7)
    add_body(doc, "共现计算、邻接构建和推荐都随线程数显著下降；归并却因线程局部哈希表数量增加而持续增长。48线程时，归并占算法时间约三分之二，完全吞噬其他阶段的并行收益。")

    add_heading(doc, "18. 平台拓扑与异常线程点", 2)
    add_body(doc, "medium Full 的4线程和16线程推荐阶段出现稳定反常；Top-50 也有较弱的同类现象。24线程后 large 明显退化，暗示平台可能是24物理核心/48逻辑线程，或存在多NUMA节点与 spread 放置影响。manifest 只记录逻辑处理器数，不能把这一推断写成确定硬件事实。")
    add_bullets(doc, [
        "建议补采 lscpu -e=CPU,CORE,SOCKET,NODE,ONLINE 与 numactl --hardware。",
        "对4/12/16/24/32/48线程比较 OMP_PROC_BIND=spread 与 close。",
        "正式论文使用中位数，并保留标准差和异常点，不删除不理想数据。",
    ])

    add_heading(doc, "19. 核心时间与端到端时间", 2)
    add_table(doc, ["large Top-50", "串行", "24线程", "32线程", "48线程"], [
        ["algorithm_ms", "34974", "9407", "9982", "10876"],
        ["end_to_end_ms", "48943", "24135", "23870", "25299"],
        ["端到端加速", "1.00×", "2.03×", "2.05×", "1.93×"],
    ], [1.45, 1.25, 1.25, 1.25, 1.3], font_size=9.0)
    add_body(doc, "large CSV 加载约14秒且保持串行，因此端到端最佳只有约2.05倍。论文应同时报告 algorithm_ms 与 end_to_end_ms：前者证明核心并行算法，后者解释 I/O 对完整应用的限制。")


def add_bucket_strategy(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "第六部分  当前策略：分桶并行归并", 1, page_break=True)
    add_heading(doc, "20. 为什么不使用全局锁", 2)
    add_body(doc, "直接让所有线程更新一个全局 PairHashMap，需要在开放寻址探测和插入过程中加锁；热门桶会产生高竞争，缓存行在核心间反复迁移，48线程下很可能比当前串行归并更慢。因此分桶策略的目标不是把现有 merge 循环简单套 parallel for，而是改变归并所有权。")

    add_heading(doc, "21. 分桶归并设计", 2)
    add_code(doc, "bucket_id = stable_hash(pair_key) % bucket_count\n\n"
                  "parallel for each bucket:\n"
                  "    create/reserve one destination map for this bucket\n"
                  "    scan thread-local maps or bucketed entries\n"
                  "    merge only keys owned by this bucket\n"
                  "    no other thread writes this destination map\n\n"
                  "after merge:\n"
                  "    iterate all bucket maps when building CSR graph")
    add_numbers(doc, [
        "选择大于线程数的桶数，例如48线程测试128/256/384桶，增加动态调度粒度。",
        "商品对使用稳定哈希映射到唯一桶，保证同一个key只在一个目标表中归并。",
        "每个桶由一个线程独占写入，因此不需要每次 pair_map_increment 加锁。",
        "分桶表保持只读后，可直接并行统计度数和构建邻接图。",
        "以现有 checksum 和逐键比较作为准入门，先在toy/small验证，再运行large。",
    ])
    add_callout(doc, "实现风险",
                "分桶本身需要扫描、分类和额外容量管理。理论估算假设分桶开销为零，真实收益一定低于上界；必须把分桶分发时间计入 merge_ms，不能只计桶内归并。",
                color=RED, fill="FCEEEE")

    add_heading(doc, "22. 基于实测的收益上界", 2)
    add_figure(doc, figures["bucket"], "图 8  large 24线程下分桶归并的理论收益区间", 6.15)
    add_body(doc, "large 24线程当前总时间9407 ms，其中归并5553 ms，非归并部分约3854 ms。若归并净加速2倍，预计总时间6631 ms、总体5.27倍；4倍时约5242 ms、6.67倍；8倍时约4548 ms、7.69倍；完全消除归并的理论上限约9.08倍。")
    add_table(doc, ["假设", "预计总时间", "预计总体加速", "相对当前"], [
        ["当前", "9407 ms", "3.72×", "基准"],
        ["归并2×", "6631 ms", "5.27×", "总时间降低29.5%"],
        ["归并4×", "5242 ms", "6.67×", "总时间降低44.3%"],
        ["归并8×", "4548 ms", "7.69×", "总时间降低51.7%"],
        ["理想消除", "3854 ms", "9.08×", "不可实现的上界"],
    ], [1.25, 1.45, 1.55, 2.25], font_size=9.0)
    add_body(doc, "考虑内存带宽与分桶分发开销，更现实的目标是归并净加速3-6倍，使 large Top-50 总加速比达到约6-7倍，并重新观察32/48线程是否恢复扩展。")

    add_heading(doc, "23. 分桶归并的审效工作流", 2)
    add_table(doc, ["阶段", "测试", "通过条件"], [
        ["B0 设计冻结", "固定桶函数、桶数和计时边界", "串行语义不变"],
        ["B1 单元测试", "同key同桶、桶覆盖、空桶、溢出", "边界全部通过"],
        ["B2 toy/small", "逐键比较全局表与分桶表", "计数和checksum完全一致"],
        ["B3 medium审效", "1/2/4/8/24线程与不同桶数", "merge下降且总时间改善"],
        ["B4 large平台", "8/12/16/24/32/48线程", "中位数改善、无内存异常"],
        ["B5 准入决策", "质量、正确性、内存、总时间", "至少一组稳定优于当前"],
    ], [1.1, 2.65, 2.75], font_size=8.8)


def add_future_and_writing(doc: Document) -> None:
    add_heading(doc, "第七部分  推荐算法扩展与论文写作建议", 1, page_break=True)
    add_heading(doc, "24. 是否需要加入时间衰减", 2)
    add_body(doc, "时间衰减主要服务推荐质量，而不是运行时间。它需要保留订单顺序或 days_since_prior_order，并将历史频次改成加权累积；计算量增加后加速比可能表面上变大，但这只是并行工作量增加，不能包装成性能优化。当前课程主线不建议把时间衰减加入默认性能基线。")
    add_body(doc, "若论文需要算法扩展，可将时间衰减做成独立开关，仅在质量实验中比较：固定 lambda，报告 Hit Rate、Precision、Recall、NDCG 和 MRR；性能实验仍使用基础评分，保证优化前后口径清晰。")

    add_heading(doc, "25. 其他后续策略的优先级", 2)
    add_table(doc, ["策略", "主要目标", "对加速比影响", "优先级"], [
        ["分桶并行归并", "消除高线程主瓶颈", "预计最大", "最高"],
        ["线程工作区复用", "减少分配器竞争", "改善绝对时间与稳定性", "中"],
        ["NUMA/绑核优化", "降低跨节点访问", "24-48线程可能明显", "高"],
        ["Top-N堆/Quickselect", "减少完整排序", "串行也变快，可能降低加速比", "中"],
        ["并行CSV加载", "改善端到端", "不改变algorithm_ms", "较低"],
        ["时间衰减/类别融合", "改善推荐质量", "通常增加计算", "可选"],
    ], [1.6, 1.8, 2.05, 1.05], font_size=8.8)

    add_heading(doc, "26. 论文中应如何组织优化故事", 2)
    add_numbers(doc, [
        "先说明为什么共现图是合理的课程模型，再给出公式和复杂度。",
        "将独立串行基线作为正确性与加速比基准，强调不能用OpenMP 1线程替代。",
        "展示基础Full方案的4.915倍和large 6.116倍，证明两个主要并行点有效。",
        "用病态归并异常说明性能审效，而不是隐藏异常。",
        "用分母预计算和Top-N说明绝对时间优化与质量门。",
        "明确解释Top-50使加速比降到1.31的Amdahl原因。",
        "用并行邻接构建恢复到3.318倍，形成针对性优化闭环。",
        "最后用1-48线程数据证明归并成为新瓶颈，并自然引出分桶策略。",
    ])
    add_callout(doc, "建议论文中心句",
                "本项目的优化过程体现了并行程序设计中的动态瓶颈迁移：算法剪枝降低了总工作量，却提高了串行阶段占比；通过分阶段计时重新定位瓶颈后，项目继续并行化邻接构建，并用48线程实验识别出归并扩展性问题。",
                color=GREEN, fill="EAF4E5")

    add_heading(doc, "27. 不能过度表述的结论", 2)
    add_bullets(doc, [
        "不能把不同机器、不同日期的绝对时间直接计算加速比。",
        "不能把 Full 串行到 Top-50 OpenMP 的组合收益称为纯 OpenMP 加速比。",
        "不能把 Hit Rate 约84%描述为工业级高精度，Precision@10约0.26、Recall@10约0.31仍有限。",
        "不能断言学校平台一定有24物理核心，除非补采lscpu/NUMA信息。",
        "分桶归并尚未实现，6-7倍属于基于阶段时间的工程估计，不是正式实验结果。",
    ])


def add_conclusion_appendix(doc: Document) -> None:
    add_heading(doc, "结论", 1, page_break=True)
    add_body(doc, "项目完成了从业务问题抽象、稀疏数据结构设计、独立串行基线、两级 OpenMP 并行到多轮性能优化的完整链条。最初 Full 方案证明了订单级共现和用户级推荐的并行价值；精确分母预计算与 Top-N 将实际运行时间显著压缩；加速比下降促使项目从“看总时间”转向“看阶段占比”，最终通过两遍式并行邻接构建恢复并行收益。")
    add_body(doc, "学校平台 48线程实验进一步证明，当前 Top-50 在24线程达到最佳，随后由于串行归并反向增长而退化。这个结果为分桶并行归并提供了直接实证依据。项目的研究价值因此不只是一条加速曲线，而是一套可复用的优化方法：保持语义、分阶段计时、提出瓶颈假设、实现最小优化、执行审效门、再用更大数据和更多线程验证。")
    add_callout(doc, "当前项目状态",
                "并行邻接构建已实现并通过正式验证；分桶并行归并处于设计与收益估算阶段。下一轮工作应以分桶正确性和large高线程审效为中心。",
                color=BLUE)

    add_heading(doc, "附录 A  指标与计时口径", 1, page_break=True)
    add_table(doc, ["字段", "含义", "论文用途"], [
        ["load_ms", "CSV与内存结构构建", "解释I/O上限"],
        ["cooccur_compute_ms", "购物篮共现局部计算", "订单级并行"],
        ["merge_ms", "线程局部统计归并", "当前主要瓶颈"],
        ["adjacency_ms", "CSR构图与Top-N", "并行邻接收益"],
        ["recommend_ms", "用户候选评分与Top-K", "用户级并行"],
        ["algorithm_ms", "共现到推荐完成", "核心加速比"],
        ["end_to_end_ms", "加载到评估结束", "完整应用体验"],
        ["Speedup", "T_serial/T_parallel", "并行收益"],
        ["Efficiency", "Speedup/p", "线程利用率"],
    ], [1.55, 2.9, 2.05], font_size=8.8)

    add_heading(doc, "附录 B  主要证据路径", 1)
    add_table(doc, ["证据", "路径"], [
        ["问题分析与数据处理", "问题分析与数据处理.docx"],
        ["串行原型说明", "串行算法实现(1).docx；basket_recommender/"],
        ["基础性能结果", "results/summary/runtime-medium-20260716-125503-summary.csv"],
        ["归并异常", "results/summary/INVALID-pre-reserve-...-summary.csv"],
        ["优化结果", "project/docs/optimization-results.md"],
        ["测试体系", "project/docs/test-report.md"],
        ["学校medium Full", "project/results/experiments/20260718-152329241-.../"],
        ["学校medium Top-50", "project/results/experiments/20260718-152600213-.../"],
        ["学校large Top-50", "project/results/experiments/20260718-154737913-.../"],
    ], [1.75, 4.75], font_size=8.5)

    add_heading(doc, "附录 C  项目优化时间线", 1)
    add_table(doc, ["阶段", "发现", "行动", "结果"], [
        ["建模", "原任务复杂", "转为共现+复购Top-K", "可解释且天然并行"],
        ["串行原型", "验证业务闭环", "C++模块化实现", "获得参考语义"],
        ["基础并行", "订单/用户独立", "线程局部统计+用户并行", "medium 4.915×"],
        ["异常修复", "OpenMP1归并12.96s", "全局表一次reserve", "归并降至48ms"],
        ["热点优化", "重复sqrt", "边分母预计算", "串行Full降35.8%"],
        ["算法剪枝", "完整邻居过多", "Top-N并选Top-50", "候选减少94.25%"],
        ["新问题", "加速比降至1.31×", "阶段计时+Amdahl分析", "定位串行构图"],
        ["并行构图", "邻接占比升高", "两遍式无锁CSR+并行Top-N", "本地3.318×"],
        ["平台验证", "24线程后退化", "1-48线程large实验", "归并占比升至66.8%"],
        ["当前策略", "串行归并主导", "分桶所有权并行归并", "待实现与审效"],
    ], [0.95, 1.75, 2.15, 1.65], font_size=8.3)


def build_document() -> Path:
    figures = create_figures()
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    set_custom_header_footer(doc)
    doc.core_properties.title = "Instacart购物篮推荐项目优化全过程深度研究"
    doc.core_properties.subject = "并行程序设计课程论文报告组参考文档"
    doc.core_properties.author = "并行程序设计项目组"
    doc.core_properties.keywords = "OpenMP, Instacart, 共现推荐, 并行邻接表, 分桶归并"

    add_cover(doc)
    add_navigation(doc)
    add_executive_summary(doc)
    add_problem_model(doc, figures)
    add_serial_parallel(doc, figures)
    add_algorithm_optimization(doc, figures)
    add_parallel_graph(doc)
    add_school_results(doc, figures)
    add_bucket_strategy(doc, figures)
    add_future_and_writing(doc)
    add_conclusion_appendix(doc)

    doc.save(OUTPUT)
    print(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    build_document()
