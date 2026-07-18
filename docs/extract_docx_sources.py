"""Extract paragraphs and tables from local DOCX files into readable Markdown."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(document: Document):
    """Yield top-level paragraphs and tables in document order."""
    body = document.element.body
    paragraph_by_element = {p._p: p for p in document.paragraphs}
    table_by_element = {t._tbl: t for t in document.tables}
    for child in body.iterchildren():
        if child in paragraph_by_element:
            yield paragraph_by_element[child]
        elif child in table_by_element:
            yield table_by_element[child]


def paragraph_markdown(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("Heading"):
        try:
            level = max(1, min(6, int(style.split()[-1])))
        except ValueError:
            level = 2
        return f"{'#' * level} {text}"
    if "Title" in style:
        return f"# {text}"
    if "List" in style:
        return f"- {text}"
    return text


def table_markdown(table: Table) -> list[str]:
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells]
            for row in table.rows]
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    result = ["| " + " | ".join(rows[0]) + " |",
              "| " + " | ".join(["---"] * width) + " |"]
    result.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return result


def extract(source: Path, target: Path) -> None:
    document = Document(source)
    lines = [f"<!-- Extracted from {source.name} -->", ""]
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = paragraph_markdown(block)
            if text:
                lines.extend((text, ""))
        else:
            lines.extend(table_markdown(block))
            lines.append("")
    target.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    args = parser.parse_args()
    args.target.parent.mkdir(parents=True, exist_ok=True)
    extract(args.source, args.target)
    print(args.target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
