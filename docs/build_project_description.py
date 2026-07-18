from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "project"
OUTPUT = ROOT / "并行程序设计大作业项目说明书.docx"
ASSET_DIR = PROJECT / "docs" / "_project_description_assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183B56"
INK = "243447"
MUTED = "657786"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "2F6B4F"
GOLD = "9A6A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C8D0D9", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_in: Sequence[float]) -> None:
    total_dxa = int(round(sum(widths_in) * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_in)):
            dxa = int(round(width * 1440))
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei",
                 size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string("263238")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.1

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        p_pr.append(borders)


def add_body(doc, text: str, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item))


def add_numbers(doc, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item))


def add_callout(doc, label: str, text: str, color=BLUE, fill=CALLOUT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, fill, color)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=color)
    set_run_font(p.add_run(text))


def add_code(doc, code: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, "F5F7F9")
    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9)
        if index != len(lines) - 1:
            run.add_break()


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[float], font_size=9.2,
              header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "FAFBFC")
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if index == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path: Path, caption: str, width=6.15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph(style="Figure Caption")
    set_run_font(cap.add_run(caption), size=9, italic=True, color=MUTED)


def add_heading(doc, text: str, level=1, page_break=False):
    p = doc.add_heading(text, level=level)
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("并行程序设计课程项目  |  项目说明书"),
                 size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("第 "), size=8.5, color=MUTED)
    add_page_number(p)
    set_run_font(p.add_run(" 页"), size=8.5, color=MUTED)


def create_architecture_diagram(path: Path) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 880), "white")
    draw = ImageDraw.Draw(image)
    regular_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    regular = ImageFont.truetype(str(regular_path), 27)
    bold = ImageFont.truetype(str(bold_path), 30)
    title_font = ImageFont.truetype(str(bold_path), 38)

    boxes = [
        (40, 180, 275, 360, "CSV data\n4 tables", "#E8EEF5"),
        (330, 180, 570, 360, "CSV Loader\nDataset CSR", "#DDEBF7"),
        (625, 180, 940, 360, "Serial / OpenMP\nCo-occurrence", "#D9EAD3"),
        (995, 180, 1260, 360, "PairHashMap\nPopularity", "#FFF2CC"),
        (1315, 180, 1750, 360, "CSR Graph\nRead-only", "#FCE5CD"),
        (350, 550, 735, 735, "Serial / OpenMP\nUser Recommendation", "#D9EAD3"),
        (795, 550, 1130, 735, "Top-K Result\nDeterministic", "#EADCF8"),
        (1190, 550, 1585, 735, "Hit Rate / Recall\nChecksums", "#DDEBF7"),
    ]

    def centered_multiline(box, text, font):
        x1, y1, x2, y2 = box
        lines = text.split("\n")
        line_height = font.size + 8
        total = line_height * len(lines) - 8
        y = (y1 + y2 - total) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            x = (x1 + x2 - (bbox[2] - bbox[0])) / 2
            draw.text((x, y), line, font=font, fill="#243447")
            y += line_height

    for x1, y1, x2, y2, label, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20,
                               fill=fill, outline="#5B6B7A", width=3)
        centered_multiline((x1, y1, x2, y2), label, bold)

    def arrow(start, end):
        draw.line((start, end), fill="#657786", width=5)
        x2, y2 = end
        x1, y1 = start
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 > x1 else -1
            tip = [(x2, y2), (x2 - 18 * direction, y2 - 12),
                   (x2 - 18 * direction, y2 + 12)]
        else:
            direction = 1 if y2 > y1 else -1
            tip = [(x2, y2), (x2 - 12, y2 - 18 * direction),
                   (x2 + 12, y2 - 18 * direction)]
        draw.polygon(tip, fill="#657786")

    arrow((275, 270), (330, 270))
    arrow((570, 270), (625, 270))
    arrow((940, 270), (995, 270))
    arrow((1260, 270), (1315, 270))
    draw.line(((1532, 360), (1532, 470), (542, 470), (542, 550)),
              fill="#657786", width=5)
    draw.polygon([(542, 550), (530, 530), (554, 530)], fill="#657786")
    arrow((735, 642), (795, 642))
    arrow((1130, 642), (1190, 642))

    title = "End-to-end architecture and data flow"
    bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((1800 - (bbox[2] - bbox[0])) / 2, 65), title,
              font=title_font, fill="#183B56")
    image.save(path, quality=95)


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(52)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("并行程序设计课程大作业"),
                 size=12, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("基于 OpenMP 的 Instacart\n购物篮推荐项目说明书"),
                 size=27, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("算法建模 · 程序架构 · 并行设计 · 审效测试 · 性能结果"),
                 size=13, color=DARK_BLUE)

    add_table(doc, ["文档属性", "内容"], [
        ["项目定位", "Kaggle 购物篮推荐任务的串行/OpenMP 并行求解与性能分析"],
        ["主要读者", "项目报告撰写组、PPT 制作组、答辩组、后续开发组员"],
        ["实现语言", "C11 + OpenMP；Python 用于性能汇总和绘图"],
        ["当前版本", "基础版本完成，G0-G10 审效测试通过"],
        ["编制日期", "2026 年 7 月 16 日"],
    ], [1.55, 4.95], font_size=10)

    add_callout(doc, "阅读建议",
                "报告撰写组优先阅读第 1、2、4、8、10、12、13 章；开发组优先阅读第 5、6、7、9、11、14 章。",
                color=GOLD, fill="FFF8E8")
    doc.add_page_break()


def build_document() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_path = ASSET_DIR / "architecture.png"
    create_architecture_diagram(architecture_path)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    doc.core_properties.title = "基于 OpenMP 的 Instacart 购物篮推荐项目说明书"
    doc.core_properties.subject = "并行程序设计课程大作业项目交接与报告写作参考"
    doc.core_properties.author = "并行程序设计课程项目组"
    doc.core_properties.keywords = "OpenMP, Instacart, 购物篮推荐, 商品共现, 并行计算"

    add_cover(doc)

    add_heading(doc, "内容导航", 1)
    add_body(doc, "本说明书不是最终课程报告，而是面向组员的统一事实底稿。报告撰写人员可以从中提取章节内容、表格数据、性能结论和局限性说明；开发人员可以据此理解模块边界、内存所有权和并行正确性约束。")
    nav = [
        "1 项目背景与目标", "2 问题建模与范围", "3 数据集与预处理",
        "4 推荐算法设计", "5 总体程序架构", "6 核心数据结构",
        "7 串行基线", "8 OpenMP 并行设计", "9 推荐与评估实现",
        "10 审效测试体系", "11 工程目录与文件职责", "12 性能实验与结果",
        "13 关键优化与结果解释", "14 构建、运行与交接", "15 报告写作指南",
        "16 局限性与扩展", "附录 A-D：命令行、输出字段、术语和验收清单",
    ]
    add_bullets(doc, nav)
    add_callout(doc, "统一口径",
                "课程报告中的加速比以独立串行实现的 algorithm_ms 为基准；OpenMP 1 线程只作为额外对照，不能代替独立串行基线。")

    add_heading(doc, "1 项目背景与目标", 1, page_break=True)
    add_heading(doc, "1.1 选题背景", 2)
    add_body(doc, "Instacart Market Basket Analysis 的原始任务是根据用户历史订单预测下一次购物篮中的商品。本项目不以 Kaggle 排名为核心，而是将该问题转化为适合并行程序设计课程展示的计算任务：在大规模历史订单上构建商品共现关系，并为用户生成 Top-K 推荐。")
    add_body(doc, "购物篮数据天然包含大量规则循环：订单聚合、购物篮内商品两两组合、商品热度统计、用户候选生成和候选评分。这些步骤既有足够的计算量，又具有良好的任务独立性，适合使用 OpenMP 讨论线程划分、负载均衡、同步、局部存储和归并开销。")
    add_heading(doc, "1.2 课程目标对应", 2)
    add_table(doc, ["课程要求", "本项目对应内容", "交付证据"], [
        ["Kaggle 问题", "Instacart 购物篮推荐", "数据集与问题建模"],
        ["算法与程序设计", "商品共现图 + Top-K 推荐", "串行/OpenMP 源码"],
        ["串并行对比", "独立串行版与 OpenMP 版", "逐项校验和与性能 CSV"],
        ["性能与效益", "时间、加速比、效率、推荐指标", "medium/large 实验结果"],
        ["报告与答辩", "架构、算法、并行优化和局限性", "本文档与图表"],
    ], [1.35, 3.25, 1.9])
    add_heading(doc, "1.3 最终目标", 2)
    add_bullets(doc, [
        "提供可人工验证的串行正确性基线。",
        "使用 OpenMP 加速商品共现统计和用户推荐两个主要计算阶段。",
        "确保不同线程数和调度方式下的结果与串行版一致。",
        "在 small、medium、large 三个规模上展示正确性、性能和扩展性。",
        "形成可复现的测试脚本、性能原始数据、汇总表和图表。",
    ])

    add_heading(doc, "2 问题建模与项目范围", 1, page_break=True)
    add_heading(doc, "2.1 从下一购物篮预测到共现推荐", 2)
    add_body(doc, "项目将每个订单视为一个购物篮，将每种商品视为图节点。如果两种商品出现在同一个 prior 购物篮中，就在二者之间建立无向加权边，边权为它们在所有历史订单中的共现次数。用户历史商品作为查询起点，其历史商品及共现邻居组成候选集。")
    add_code(doc, "用户历史商品\n    -> 查询共现邻居\n    -> 累计个人频次、归一化共现和全局热度\n    -> 确定性排序\n    -> Top-K 推荐")
    add_heading(doc, "2.2 建模合理性", 2)
    add_bullets(doc, [
        "算法直观：小规模订单可以人工枚举商品对并核对计数。",
        "计算密集：长度为 m 的购物篮产生 m(m-1)/2 个商品位置对。",
        "数据独立：不同订单的局部共现计算互不依赖。",
        "并行问题典型：存在负载不均、线程局部表、归并串行段和内存带宽限制。",
        "课程重点明确：推荐精度用于验证实用性，主要分析对象仍是并行性能。",
    ])
    add_heading(doc, "2.3 基础版本边界", 2)
    add_table(doc, ["纳入基础版本", "暂不纳入基础版本"], [
        ["商品共现、商品热度、用户历史频次", "时间衰减和购买周期模型"],
        ["串行/OpenMP 共现和推荐", "MPI、pthread 完整版本"],
        ["Hit Rate@K、宏平均 Recall@K", "负采样、AUC、复杂机器学习模型"],
        ["static/dynamic 调度对比", "并行排序和分布式图计算"],
        ["完整共现图", "强制 Top-N 邻居截断"],
    ], [3.25, 3.25])
    add_callout(doc, "范围说明",
                "large 完整全图已在 27.81 GiB 机器上成功运行，因此 Top-N 截断不是基础版本可运行性的必要条件，只保留为后续内存优化选项。")

    add_heading(doc, "3 数据集与预处理", 1, page_break=True)
    add_heading(doc, "3.1 输入文件", 2)
    add_table(doc, ["文件", "主要字段", "程序用途"], [
        ["orders.csv", "order_id, user_id, eval_set, order_number", "订单索引、用户映射、prior/train 区分"],
        ["order_products__prior.csv", "order_id, product_id, add_to_cart_order", "历史购物篮、用户历史、共现统计"],
        ["order_products__train.csv", "order_id, product_id", "下一购物篮验证真值"],
        ["products.csv", "product_id, product_name, aisle_id, department_id", "有效商品集合和最大商品 ID"],
    ], [1.75, 2.65, 2.1], font_size=8.8)
    add_heading(doc, "3.2 数据规模", 2)
    add_table(doc, ["规模", "订单数", "prior 明细", "train 明细", "商品数", "用途"], [
        ["toy", "7", "6", "4", "4", "手算测试"],
        ["small", "2,995", "26,575", "1,327", "6,170", "正确性回归"],
        ["medium", "81,832", "761,750", "33,782", "29,330", "重复性能实验"],
        ["large", "3,421,083", "32,434,489", "1,384,617", "49,685", "完整规模验证"],
    ], [0.7, 1.0, 1.15, 1.05, 0.9, 1.7], font_size=8.6)
    add_heading(doc, "3.3 加载和结构化过程", 2)
    add_numbers(doc, [
        "读取 products.csv，建立有效商品标记和最大商品 ID。",
        "读取 orders.csv，建立 order_id 到 user_id、eval_set 和 prior 购物篮下标的稠密索引。",
        "两遍读取 prior 明细：第一遍统计每个购物篮长度，第二遍一次性填充连续商品数组。",
        "把 prior 的 (user_id, product_id) 记录排序压缩，得到用户历史频次 CSR。",
        "把 train 的 (user_id, product_id) 记录排序去重，得到用户验证真值 CSR。",
        "执行 offset 单调性、末 offset、ID 范围和明细数量等结构不变量检查。",
    ])
    add_callout(doc, "为什么两遍读取",
                "第一遍得到准确长度后再一次性分配，可避免 large 数据上频繁 realloc 和大量小内存块；代价是顺序读取 prior 文件两次。")

    add_heading(doc, "4 推荐算法设计", 1, page_break=True)
    add_heading(doc, "4.1 商品共现统计", 2)
    add_body(doc, "对每个 prior 购物篮枚举所有 i<j 的商品位置对。商品对按 min(product_i, product_j) 和 max(product_i, product_j) 规范化，使 (a,b) 与 (b,a) 使用同一个键。")
    add_code(doc, "uint32_t a = min(product_i, product_j);\nuint32_t b = max(product_i, product_j);\nuint64_t key = ((uint64_t)a << 32) | b;\npair_count[key] += 1;")
    add_body(doc, "若第 r 个购物篮长度为 m_r，商品对枚举总复杂度为 Σ m_r(m_r-1)/2。空间复杂度与实际出现的唯一商品对数量相关，而不是商品种类数平方。")
    add_heading(doc, "4.2 候选集", 2)
    add_bullets(doc, [
        "用户历史购买过的商品，用于复购预测。",
        "用户历史商品在共现图中的全部邻居，用于关联推荐。",
        "同一候选商品只保留一个分数槽位，来自多个历史商品的贡献累加。",
    ])
    add_heading(doc, "4.3 评分函数", 2)
    add_code(doc, "score(u,p) = 1.0 * freq(u,p)\n           + 0.8 * co_score(u,p)\n           + 0.2 * log(1 + popularity(p))\n\nco_score(u,p) = sum_q freq(u,q) * cooccur(q,p)\n                / sqrt(popularity(q) * popularity(p))")
    add_body(doc, "归一化共现项降低纯热门商品的支配作用；log(1+popularity) 让全局热度成为弱补充；用户历史频次直接支持 Instacart 场景中的复购。参数固定是为了把实验重点放在并行性能，而不是参数调优。")
    add_heading(doc, "4.4 确定性 Top-K", 2)
    add_body(doc, "候选按分数降序排列；若分数完全相同，则商品 ID 小者优先。程序使用 K 个槽位的插入式维护，K=10 时每个候选的排序成本近似常数。确定性次级规则是串并行结果逐项比较的基础。")

    add_heading(doc, "5 总体程序架构", 1, page_break=True)
    add_figure(doc, architecture_path, "图 5-1  项目端到端架构与数据流")
    add_heading(doc, "5.1 分层职责", 2)
    add_table(doc, ["层次", "主要结构/模块", "职责"], [
        ["输入层", "csv_loader", "读取、校验并构建连续内存结构"],
        ["模型层", "PairHashMap, CooccurResult", "统计商品对和商品热度"],
        ["图层", "CooccurGraph", "把无序商品对转换为双向 CSR 邻接表"],
        ["推荐层", "recommender", "候选生成、评分和确定性 Top-K"],
        ["评估层", "evaluator", "Hit Rate、Recall 和校验和"],
        ["实验层", "PowerShell/Python scripts", "批量运行、汇总和绘图"],
    ], [1.1, 2.1, 3.3])
    add_heading(doc, "5.2 模块解耦", 2)
    add_body(doc, "串行和 OpenMP 实现共享相同 Dataset、CooccurResult、CooccurGraph 和 RecommendationResult 接口。这样既保证两种版本做相同工作，也让测试程序可以直接执行逐项比较，而不是只比较最终推荐指标。")
    add_callout(doc, "架构关键点",
                "并行版本不是对串行代码简单加 pragma：它显式引入线程局部哈希表、线程局部热度数组、归并阶段和线程私有推荐工作区。")

    add_heading(doc, "6 核心数据结构与内存设计", 1, page_break=True)
    add_heading(doc, "6.1 CSR 连续存储", 2)
    add_body(doc, "BasketTable、UserHistory、GroundTruth 和 CooccurGraph 均使用类似 CSR 的 offsets + values 结构。第 i 组数据占用 [offsets[i], offsets[i+1])。这种设计减少指针和内存碎片，提升顺序访问局部性，并允许 OpenMP 线程安全地只读共享。")
    add_table(doc, ["结构", "offsets 的索引", "values 内容"], [
        ["BasketTable", "prior 购物篮下标", "product_id"],
        ["UserHistory", "user_id", "product_id + frequency"],
        ["GroundTruth", "user_id", "train product_id"],
        ["CooccurGraph", "product_id", "neighbor_id + weight"],
    ], [1.4, 2.1, 3.0])
    add_heading(doc, "6.2 开放寻址商品对哈希表", 2)
    add_body(doc, "PairHashMap 的每个槽位保存 64 位 key、32 位 count 和 used 标记。哈希冲突采用线性探测；负载因子接近 0.70 时扩容。归并前通过 pair_map_reserve 按线程局部键数上界预留容量，避免批量插入形成病态长探测链。")
    add_heading(doc, "6.3 推荐工作区", 2)
    add_body(doc, "C 版本没有为每个用户创建 unordered_map，而是每线程复用 scores、marks 和 candidates 三个数组。generation 每处理一个用户递增，marks[product_id]==generation 表示该商品属于当前用户候选。这样无需为每个用户清空约 5 万个分数位置。")
    add_heading(doc, "6.4 内存所有权", 2)
    add_table(doc, ["创建函数", "释放函数", "所有权说明"], [
        ["dataset_load", "dataset_free", "调用方持有完整输入数据"],
        ["build_cooccur_*", "cooccur_result_free", "调用方持有共现表和热度数组"],
        ["cooccur_graph_build", "cooccur_graph_free", "调用方持有 CSR 图"],
        ["recommend_*", "recommendation_result_free", "调用方持有用户 Top-K"],
    ], [2.1, 2.0, 2.4])

    add_heading(doc, "7 独立串行基线", 1, page_break=True)
    add_body(doc, "串行版用于提供正确性和性能基准，不能由 OpenMP 版本设置 1 个线程替代。二者虽然算法等价，但 OpenMP 版本仍包含线程局部结构和归并成本。")
    add_numbers(doc, [
        "串行遍历 prior 购物篮，统计商品热度。",
        "串行枚举每个购物篮的无序商品对并更新 PairHashMap。",
        "把全局商品对表转换为双向 CSR 共现图。",
        "串行遍历用户，生成候选、计算分数并维护 Top-K。",
        "串行评估推荐结果并输出稳定校验和。",
    ])
    add_heading(doc, "7.1 串行版的重要性", 2)
    add_bullets(doc, [
        "toy 数据可直接人工核对。",
        "OpenMP 任意线程数必须与它逐项一致。",
        "加速比 T_serial/T_parallel 使用它的 algorithm_ms。",
        "能够识别 OpenMP 1 线程的额外管理和归并开销。",
    ])

    add_heading(doc, "8 OpenMP 并行设计", 1, page_break=True)
    add_heading(doc, "8.1 订单级共现", 2)
    add_code(doc, "#pragma omp parallel\n{\n    int tid = omp_get_thread_num();\n    PairHashMap *pairs = &local_pairs[tid];\n    uint32_t *pop = local_pop + tid * product_count;\n\n    #pragma omp for schedule(dynamic, 64)\n    for each basket {\n        update thread-local popularity;\n        enumerate pairs into thread-local hash map;\n    }\n}")
    add_body(doc, "不同订单之间独立，但购物篮长度差异会使计算量按平方变化，因此默认使用 dynamic,64。线程只写自己的局部表和局部热度数组，输入购物篮只读。")
    add_heading(doc, "8.2 归并阶段", 2)
    add_body(doc, "并行循环结束后，基础版串行合并线程局部表。该设计没有锁且易于验证，但属于 Amdahl 定律中的串行部分。程序将 cooccur_compute_ms 与 merge_ms 分开记录，以便判断是否需要进一步并行归并。")
    add_heading(doc, "8.3 用户级推荐", 2)
    add_code(doc, "#pragma omp parallel for schedule(dynamic, 16)\nfor (user = 0; user < user_count; ++user) {\n    use thread-private workspace;\n    build candidates and scores;\n    write recommendations[user * K ... (user + 1) * K);\n}")
    add_body(doc, "每个用户写独立结果区间，模型、历史和共现图只读，因此不需要 critical 或 atomic。用户历史长度和邻居数量不同，dynamic,16 有助于平衡工作量。")
    add_heading(doc, "8.4 线程安全分析", 2)
    add_table(doc, ["数据", "访问方式", "线程安全依据"], [
        ["Dataset / CooccurGraph", "共享只读", "构建完成后不再修改"],
        ["local_pairs[tid]", "线程私有写", "仅 tid 线程访问"],
        ["local_pop[tid]", "线程私有写", "按线程分片"],
        ["RecommendationWorkspace", "每线程一份", "不跨线程共享"],
        ["recommendations[u]", "用户独占写", "每个循环迭代写不同区间"],
        ["全局 PairHashMap", "归并期串行写", "并行区结束后访问"],
    ], [1.7, 1.55, 3.25])

    add_heading(doc, "9 推荐评估与结果一致性", 1, page_break=True)
    add_heading(doc, "9.1 Hit Rate@K", 2)
    add_body(doc, "对每个有 train 真值的用户，只要 Top-K 中至少命中一个真实商品，该用户记为命中。Hit Rate 是命中用户数除以评估用户数。")
    add_heading(doc, "9.2 Recall@K", 2)
    add_body(doc, "先计算每个用户被命中的真实商品数占该用户 train 商品数的比例，再对所有有 train 真值的用户求宏平均。空真值用户不进入分母。")
    add_heading(doc, "9.3 校验和", 2)
    add_body(doc, "程序同时输出 cooccur_checksum 和 recommendation_checksum。正式实验中，不同线程数、调度方式和串行/OpenMP 模式必须具有相同校验和。校验和用于快速回归；调试阶段仍执行完整逐项比较。")
    add_heading(doc, "9.4 当前推荐效果", 2)
    add_table(doc, ["数据规模", "Hit Rate@10", "Recall@10", "说明"], [
        ["medium", "0.8428125", "0.314722676980", "所有线程数一致"],
        ["large", "0.843189110503", "0.306467486678", "串行与 8 线程一致"],
    ], [1.25, 1.35, 1.35, 2.55])
    add_callout(doc, "效果解释",
                "推荐效果不是本课程的主要优化目标。当前指标证明基础共现模型具有实际预测能力，同时不会因并行化改变算法语义。")

    add_heading(doc, "10 审效测试体系", 1, page_break=True)
    add_body(doc, "项目使用“实现 - 单元测试 - 串并行比较 - 性能有效性审查”的质量门。任何并行优化都必须重新通过相关正确性测试。")
    add_table(doc, ["质量门", "核心检查", "结果"], [
        ["G0", "GCC/OpenMP、数据表头、toy 数据", "通过"],
        ["G1", "加载、CSR、重复释放、错误路径", "通过"],
        ["G2", "哈希冲突、扩容、reserve、溢出、随机预言机", "通过"],
        ["G3", "toy 串行共现手算", "通过"],
        ["G4", "1/2/4 线程及 static/dynamic 与串行逐项一致", "通过"],
        ["G5", "图双向性、度数和、offset、排序", "通过"],
        ["G6", "评分、Top-K、空历史、同分规则", "通过"],
        ["G7", "串并行用户推荐逐项一致", "通过"],
        ["G8", "Hit Rate/Recall 和端到端回归", "通过"],
        ["G9", "Release 重复性能实验和计时口径", "通过"],
        ["G10", "large 完整规模和最终交付", "通过"],
    ], [0.65, 4.95, 0.9], font_size=8.6)
    add_heading(doc, "10.1 toy 手算事实", 2)
    add_code(doc, "Order 10: [1,2,3]\nOrder 11: [1,2]\nOrder 12: [2]\nOrder 13: []\n\nPair counts: (1,2)=2, (1,3)=1, (2,3)=1\nPopularity: 1=2, 2=3, 3=1, 4=0\nTotal pair events: 4")
    add_heading(doc, "10.2 完整回归命令", 2)
    add_code(doc, ".\\scripts\\run_correctness.ps1")
    add_body(doc, "看到 PASS: complete correctness regression 表示加载、哈希、共现、推荐、评估和端到端校验和全部通过。")

    add_heading(doc, "11 工程目录与文件职责", 1, page_break=True)
    add_heading(doc, "11.1 头文件", 2)
    add_table(doc, ["文件", "职责"], [
        ["include/model.h", "核心数据结构和 CSR 图定义"],
        ["include/csv_loader.h", "数据加载、验证和释放接口"],
        ["include/pair_hash.h", "商品对哈希表接口"],
        ["include/cooccurrence.h", "串行/OpenMP 共现与图接口"],
        ["include/recommender.h", "串行/OpenMP 推荐接口"],
        ["include/evaluator.h", "推荐指标接口"],
    ], [2.1, 4.4])
    add_heading(doc, "11.2 源文件", 2)
    add_table(doc, ["文件", "职责"], [
        ["src/main.c", "命令行解析和端到端流程"],
        ["src/csv_loader.c", "四类 CSV 和三类 CSR 数据"],
        ["src/pair_hash.c", "开放寻址哈希表"],
        ["src/cooccurrence_serial.c", "串行共现、结果比较和图构建"],
        ["src/cooccurrence_openmp.c", "订单并行、局部表和归并"],
        ["src/recommender_serial.c", "评分、Top-K 和串行推荐"],
        ["src/recommender_openmp.c", "用户级并行推荐"],
        ["src/evaluator.c", "Hit Rate 和 Recall"],
    ], [2.25, 4.25])
    add_heading(doc, "11.3 脚本和文档", 2)
    add_table(doc, ["文件", "用途"], [
        ["scripts/build.ps1", "Debug/Release/测试目标构建"],
        ["scripts/run_correctness.ps1", "完整正确性回归"],
        ["scripts/run_benchmark.ps1", "重复性能实验和原始 CSV"],
        ["scripts/summarize_results.py", "中位数、加速比和效率"],
        ["scripts/plot_results.py", "四类性能图"],
        ["docs/test-report.md", "G0-G10 测试证据"],
        ["docs/performance-notes.md", "性能结论和优化记录"],
        ["docs/known-issues.md", "环境限制和已知问题"],
    ], [2.45, 4.05])

    add_heading(doc, "12 性能实验设计与结果", 1, page_break=True)
    add_heading(doc, "12.1 实验环境与计时边界", 2)
    add_table(doc, ["项目", "配置"], [
        ["CPU", "AMD Ryzen 7 8845H，8 核 16 线程"],
        ["内存", "约 27.81 GiB"],
        ["系统", "Windows NT 10.0.26200.0"],
        ["编译器", "MinGW-W64 GCC 8.1.0"],
        ["Release 参数", "-O2 -DNDEBUG -std=c11 -fopenmp"],
        ["线程数", "1、2、4、8"],
        ["重复次数", "medium 每组 3 次，取中位数"],
    ], [1.65, 4.85])
    add_body(doc, "algorithm_ms 包含共现计算、线程局部归并、邻接图构建和推荐，不包含 CSV I/O、指标评估和正确性比较。end_to_end_ms 包含加载到评估的完整流程。")
    add_heading(doc, "12.2 medium 正式结果", 2)
    add_table(doc, ["版本", "线程", "algorithm_ms", "Speedup", "Efficiency"], [
        ["serial", "1", "7,711", "1.000", "100.0%"],
        ["OpenMP", "1", "7,885", "0.978", "97.8%"],
        ["OpenMP", "2", "4,224", "1.826", "91.3%"],
        ["OpenMP", "4", "2,668", "2.890", "72.3%"],
        ["OpenMP", "8", "1,569", "4.915", "61.4%"],
    ], [1.25, 0.85, 1.65, 1.35, 1.4])
    figures = PROJECT / "results" / "figures"
    add_figure(doc, figures / "medium-runtime.png", "图 12-1  medium 不同线程数算法时间")
    add_figure(doc, figures / "medium-speedup.png", "图 12-2  medium OpenMP 加速比与理想加速比")
    add_figure(doc, figures / "medium-efficiency.png", "图 12-3  medium 并行效率")
    add_figure(doc, figures / "medium-stages.png", "图 12-4  medium 各阶段耗时构成")
    add_heading(doc, "12.3 large 完整规模", 2)
    add_table(doc, ["版本", "线程", "algorithm_ms", "end_to_end_ms", "Speedup", "Efficiency"], [
        ["serial", "1", "1,217,953", "1,234,896", "1.000", "100.0%"],
        ["OpenMP", "8", "199,149", "217,511", "6.116", "76.4%"],
    ], [1.0, 0.7, 1.35, 1.35, 1.0, 1.1])
    add_body(doc, "large 共产生 40,749,010 条唯一商品对和 238,428,378 次商品对事件。串行与 OpenMP 的推荐指标、共现校验和和推荐校验和完全一致。")

    add_heading(doc, "13 关键优化与性能结果解释", 1, page_break=True)
    add_heading(doc, "13.1 线程局部统计", 2)
    add_body(doc, "如果所有线程直接更新全局哈希表，就需要对高频插入加锁，临界区会成为主要瓶颈。线程局部表将同步从每个商品对事件推迟到并行循环结束后的批量归并。")
    add_heading(doc, "13.2 病态归并问题与修复", 2)
    add_body(doc, "初次 medium 实验中，OpenMP 1 线程归并约 12.96 秒，而 2-8 线程仅约 0.1 秒。结果正确但性能明显异常。原因是全局哈希表从 1,024 槽开始，按局部哈希槽顺序批量插入，早期键在小表中形成长线性探测链。")
    add_table(doc, ["阶段", "OpenMP 1 线程 merge_ms"], [
        ["修复前", "约 12,960 ms"],
        ["修复后", "约 48 ms"],
    ], [3.8, 2.7])
    add_body(doc, "修复方法是在归并前统计线程局部键数上界，并调用 pair_map_reserve 一次性预留全局容量。修复后重新执行完整正确性回归和正式性能实验。")
    add_heading(doc, "13.3 加速比为何不是线性", 2)
    add_bullets(doc, [
        "邻接图构建和局部表归并仍是串行阶段。",
        "线程数增加后内存带宽和缓存竞争上升。",
        "动态调度和 OpenMP 运行时有固定开销。",
        "不同用户和购物篮的任务量并不完全均衡。",
        "8 线程效率下降不代表并行失败，而是 Amdahl 定律和硬件资源限制的正常体现。",
    ])
    add_heading(doc, "13.4 static 与 dynamic", 2)
    add_body(doc, "medium 8 线程 static 单次控制实验 algorithm_ms=1,689，dynamic 三次中位数为 1,569。当前数据上 dynamic 略优，但二者结果完全一致。报告应把该比较解释为负载均衡和调度开销的权衡，而不是绝对结论。")

    add_heading(doc, "14 构建、运行与组员交接", 1, page_break=True)
    add_heading(doc, "14.1 构建", 2)
    add_code(doc, "cd project\n.\\scripts\\build.ps1 Debug\n.\\scripts\\build.ps1 Release\n.\\scripts\\build.ps1 Smoke\n.\\build\\omp_smoke.exe 4")
    add_heading(doc, "14.2 串行和 OpenMP 运行", 2)
    add_code(doc, ".\\build\\basket_recommender.exe --data data\\medium --mode serial --top-k 10\n\n.\\build\\basket_recommender.exe --data data\\medium --mode openmp --threads 8 --top-k 10 --schedule dynamic")
    add_heading(doc, "14.3 推荐样例和文件输出", 2)
    add_code(doc, ".\\build\\basket_recommender_debug.exe --data data\\toy --mode serial --samples 3 --output build\\toy-summary.txt")
    add_heading(doc, "14.4 接手顺序", 2)
    add_numbers(doc, [
        "阅读 README 和本说明书，理解术语与数据流。",
        "阅读 include/model.h 和 src/main.c，掌握结构与端到端顺序。",
        "阅读串行共现和串行推荐，建立正确性基准认知。",
        "阅读 OpenMP 文件，重点理解线程局部数据和写区间。",
        "运行 scripts/run_correctness.ps1，确认本机环境和代码一致。",
        "在 medium 上运行 Release OpenMP 示例，不要直接从 large 开始调试。",
    ])

    add_heading(doc, "15 项目报告写作指南", 1, page_break=True)
    add_heading(doc, "15.1 建议报告结构", 2)
    add_table(doc, ["报告章节", "可引用的项目事实", "建议图表"], [
        ["背景与题目", "Instacart 下一购物篮预测；课程重点是并行性能", "问题流程图"],
        ["数据处理", "4 个 CSV、三种规模、CSR 构建和一致性检查", "数据规模表"],
        ["问题建模", "无向加权共现图、候选集和评分函数", "建模流程"],
        ["串行算法", "独立基线、复杂度和 Top-K", "串行流程图"],
        ["并行设计", "订单级和用户级并行、线程局部数据、动态调度", "架构图"],
        ["测试验证", "toy 手算、串并行逐项比较、G0-G10", "测试矩阵"],
        ["性能实验", "medium 中位数、large 完整规模、加速比和效率", "四类性能图"],
        ["优化分析", "全局预留修复病态归并、Amdahl 定律", "修复前后表"],
        ["不足与展望", "时间、类别、Top-N、并行归并、MPI/pthread", "扩展清单"],
    ], [1.25, 3.45, 1.8], font_size=8.4)
    add_heading(doc, "15.2 可以直接使用的核心结论", 2)
    add_bullets(doc, [
        "订单和用户是两个天然并行粒度，且并行语义清晰。",
        "线程局部哈希表消除了共现内层循环的全局锁竞争。",
        "medium 8 线程加速 4.915 倍，large 8 线程加速 6.116 倍。",
        "所有线程数的推荐指标和校验和一致，证明并行化没有改变结果。",
        "推荐阶段是 large 的主要瓶颈，也是用户级并行获得显著收益的原因。",
        "归并预留优化说明数据结构和插入顺序会显著影响并行程序实际性能。",
    ])
    add_heading(doc, "15.3 写作时必须保留的限定条件", 2)
    add_bullets(doc, [
        "medium 数据为每组 3 次中位数；large 为串行和 8 线程各一次完整验证。",
        "核心加速比使用 algorithm_ms，不包含 CSV I/O。",
        "推荐模型是基础共现模型，不能宣称达到 Kaggle 最优精度。",
        "当前未记录 large 的精确峰值工作集，只能说明在 27.81 GiB 机器上成功运行。",
        "static/dynamic 比较中 static 只有单次控制实验，不宜做过度统计推断。",
    ])
    add_callout(doc, "避免错误表述",
                "不要写“8 线程效率为 614%”。正确说法是 large 加速比 6.116、并行效率约 76.4%；medium 加速比 4.915、效率约 61.4%。",
                color=RED, fill="FDECEC")

    add_heading(doc, "16 局限性与后续扩展", 1, page_break=True)
    add_table(doc, ["局限性", "当前影响", "可选扩展"], [
        ["无时间衰减", "远期订单和近期订单权重相同", "按 order_number 或时间间隔加权"],
        ["未利用类别", "aisle/department 信息闲置", "增加类别亲和度项"],
        ["冷启动", "新用户/新商品只能依赖热度", "热门商品或类别回退"],
        ["完整共现图较大", "large 内存和图构建成本高", "Top-N 邻居截断"],
        ["归并仍串行", "线程数增加后占比上升", "按 key 分桶并行归并"],
        ["只有 OpenMP", "未比较多进程/手工线程", "MPI 或 pthread 对照版本"],
        ["指标较基础", "不能全面反映排序质量", "Precision、F1、NDCG 或 AUC"],
    ], [1.3, 2.5, 2.7], font_size=8.7)
    add_body(doc, "扩展应在基础版本保持可复现的前提下进行。每次优化只改变一个主要因素，保留优化前基线，重新执行完整正确性回归，并用相同数据和线程配置重新测量。")

    add_heading(doc, "17 总结", 1, page_break=True)
    add_body(doc, "本项目把 Instacart 下一购物篮预测转化为商品共现图和用户 Top-K 推荐问题，在 C11 中实现了可验证的独立串行基线，并使用 OpenMP 对订单级共现统计和用户级推荐进行并行化。工程上通过 CSR 连续存储、线程局部哈希表、全局容量预留和线程私有推荐工作区控制内存访问与同步开销。")
    add_body(doc, "测试方面，toy 提供手算基准，small 承担快速回归，medium 用于多线程重复实验，large 验证完整数据的可扩展性。所有串并行结果保持一致。性能方面，medium 8 线程加速 4.915 倍，large 8 线程加速 6.116 倍，证明两个主要并行点具有实际收益。")
    add_callout(doc, "项目交付状态",
                "基础源码、测试脚本、性能原始数据、汇总结果、图表、README、测试报告和本项目说明书均已形成，可直接供报告撰写、PPT 制作和答辩准备使用。",
                color=GREEN, fill="EAF4EE")

    add_heading(doc, "附录 A 命令行参数", 1, page_break=True)
    add_table(doc, ["参数", "含义", "默认值"], [
        ["--data DIR", "包含 4 个 CSV 的数据目录", "data/toy"],
        ["--mode serial|openmp", "运行模式", "serial"],
        ["--threads N", "OpenMP 线程数", "运行时最大线程数"],
        ["--top-k K", "每个用户推荐数量", "10"],
        ["--schedule static|dynamic", "共现循环调度", "dynamic"],
        ["--samples N", "输出前 N 个非空用户推荐样例", "不输出"],
        ["--output FILE", "将 key=value 摘要写入文件", "标准输出"],
    ], [2.0, 3.3, 1.2])

    add_heading(doc, "附录 B 主要输出字段", 1)
    add_table(doc, ["字段", "含义"], [
        ["unique_pairs", "唯一无序商品对数量"],
        ["pair_events", "全部购物篮的商品对事件总数"],
        ["load_ms", "CSV 和内存结构加载时间"],
        ["cooccur_compute_ms", "共现局部计算时间"],
        ["merge_ms", "OpenMP 局部结果归并时间"],
        ["adjacency_ms", "CSR 图构建时间"],
        ["recommend_ms", "用户推荐时间"],
        ["algorithm_ms", "共现、归并、图和推荐总时间"],
        ["end_to_end_ms", "加载到评估的完整时间"],
        ["hit_rate_at_10 / recall_at_10", "推荐效果指标"],
        ["cooccur_checksum", "共现结果校验和"],
        ["recommendation_checksum", "推荐结果校验和"],
    ], [2.55, 3.95])

    add_heading(doc, "附录 C 术语表", 1, page_break=True)
    add_table(doc, ["术语", "说明"], [
        ["购物篮", "同一个订单中的商品集合"],
        ["共现", "两种商品在同一购物篮中同时出现"],
        ["CSR", "用 offsets 和连续 values 表示变长分组/邻接表"],
        ["Top-K", "按分数返回前 K 个商品"],
        ["线程局部", "每个线程独立持有、其他线程不写的数据"],
        ["归并", "把线程局部统计累加到全局结果"],
        ["动态调度", "OpenMP 运行时按块分配剩余循环迭代"],
        ["加速比", "串行时间除以 p 线程并行时间"],
        ["并行效率", "加速比除以线程数 p"],
        ["校验和", "对大规模结果生成的稳定摘要，用于快速一致性检查"],
    ], [1.65, 4.85])

    add_heading(doc, "附录 D 交付前验收清单", 1)
    add_bullets(doc, [
        "Release 和 Debug 均能构建，编译无未处理警告。",
        "scripts/run_correctness.ps1 输出完整 PASS。",
        "报告引用的是正式 medium CSV，不是 INVALID-pre-reserve 文件。",
        "串行与并行共现/推荐校验和一致。",
        "图表中的线程数、加速比、效率和正文数字一致。",
        "明确 medium 为三次中位数，large 为单次完整验证。",
        "报告没有把 OpenMP 1 线程当作独立串行基线。",
        "报告如实说明精确峰值内存未采集。",
        "源码修改后重新执行正确性回归。",
        "large 不用于日常调试，避免重复耗费约 20 分钟串行时间。",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
